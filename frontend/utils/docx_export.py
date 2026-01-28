"""
DOCX Export Module for EasyRead

Provides functionality to export Easy Read documents to Microsoft Word format.
"""

import os
import tempfile
from typing import Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def export_to_docx(
    title: str,
    sentences: list[dict],
    image_data: Optional[dict[str, bytes]] = None,
    output_path: Optional[str] = None,
) -> str:
    """
    Export Easy Read content to a Word document.

    Args:
        title: Document title
        sentences: List of sentence dictionaries with keys:
            - sentence: The text content
            - image_prompt: Description of the image
            - highlighted: Whether this is a key sentence
            - image_path: Optional path to the icon image
        image_data: Optional dict mapping image paths to image bytes
        output_path: Optional output file path. If None, creates temp file.

    Returns:
        Path to the generated DOCX file
    """
    doc = Document()

    # Set document title
    title_paragraph = doc.add_heading(title, level=0)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add spacing after title
    doc.add_paragraph()

    # Create a table for side-by-side layout (image | text)
    for item in sentences:
        sentence = item.get("sentence", "")
        highlighted = item.get("highlighted", False)
        image_path = item.get("image_path", "")

        # Create a 1-row, 2-column table for each sentence
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        # Set column widths
        for cell in table.columns[0].cells:
            cell.width = Inches(1.5)
        for cell in table.columns[1].cells:
            cell.width = Inches(5.0)

        row = table.rows[0]
        icon_cell = row.cells[0]
        text_cell = row.cells[1]

        # Add image to the icon cell if available
        image_key = image_path.split("/")[-1] if image_path else None
        if image_data and image_key and image_key in image_data:
            img_bytes = image_data[image_key]
            # Save image to temp file for docx insertion
            with tempfile.NamedTemporaryFile(
                suffix=".png", delete=False
            ) as tmp_img:
                tmp_img.write(img_bytes)
                tmp_img_path = tmp_img.name

            try:
                icon_paragraph = icon_cell.paragraphs[0]
                icon_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = icon_paragraph.add_run()
                run.add_picture(tmp_img_path, width=Inches(1.2))
            finally:
                # Clean up temp file
                if os.path.exists(tmp_img_path):
                    os.remove(tmp_img_path)
        else:
            # Placeholder for missing image
            icon_paragraph = icon_cell.paragraphs[0]
            icon_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = icon_paragraph.add_run("[Icon]")
            run.font.color.rgb = RGBColor(128, 128, 128)

        # Add text to the text cell
        text_paragraph = text_cell.paragraphs[0]
        text_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        run = text_paragraph.add_run(sentence)
        run.font.size = Pt(14)

        # Highlight key sentences
        if highlighted:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 102, 204)

        # Add some spacing between rows
        doc.add_paragraph()

    # Generate output path if not provided
    if not output_path:
        with tempfile.NamedTemporaryFile(
            suffix=".docx", prefix="easyread_", delete=False
        ) as tmp:
            output_path = tmp.name

    doc.save(output_path)
    return output_path


def export_to_docx_bytes(
    title: str,
    sentences: list[dict],
    image_data: Optional[dict[str, bytes]] = None,
) -> bytes:
    """
    Export Easy Read content to a Word document and return as bytes.

    Args:
        title: Document title
        sentences: List of sentence dictionaries
        image_data: Optional dict mapping image paths to image bytes

    Returns:
        DOCX file content as bytes
    """
    # Create temp file, export, read, and clean up
    output_path = export_to_docx(title, sentences, image_data)

    try:
        with open(output_path, "rb") as f:
            return f.read()
    finally:
        if os.path.exists(output_path):
            os.remove(output_path)
