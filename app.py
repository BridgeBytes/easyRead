"""
EasyRead - HuggingFace Space

A single-file Gradio application combining all backend and frontend logic
for the EasyRead document simplification service.

Required environment variable (set as HF Space secret):
    GOOGLE_API_KEY  — Your Google Gemini API key
                      (GEMINI_API_KEY is also accepted as an alias)

GPU:
    Set the Space hardware to GPU (T4 or better) for fast icon generation.
    Falls back to CPU automatically if no GPU is available.
"""

import os
import json
import base64
import logging
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Optional
from uuid import uuid4

import requests
import torch
import yaml
import gradio as gr
from PIL import Image
from google import genai
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ──────────────────────────────────────────────────────────────────────────────
# Logging
# ──────────────────────────────────────────────────────────────────────────────

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
)
logger = logging.getLogger(__name__)

# ──────────────────────────────────────────────────────────────────────────────
# Paths
# ──────────────────────────────────────────────────────────────────────────────

BASE_DIR = Path(__file__).parent
CONFIG_DIR = BASE_DIR / "config"
WEIGHTS_DIR = BASE_DIR / "weights"
TEMP_DIR = Path(tempfile.gettempdir()) / "easyread"
TEMP_DIR.mkdir(parents=True, exist_ok=True)

# ──────────────────────────────────────────────────────────────────────────────
# Device detection  (GPU-first for HF Spaces)
# ──────────────────────────────────────────────────────────────────────────────

if torch.cuda.is_available():
    DEVICE = "cuda"
    TORCH_DTYPE = torch.float16
    logger.info("CUDA GPU detected — using float16 for fast inference")
elif torch.backends.mps.is_available():
    DEVICE = "mps"
    TORCH_DTYPE = torch.float32
    logger.info("Apple MPS detected")
else:
    DEVICE = "cpu"
    TORCH_DTYPE = torch.float32
    logger.info("No GPU found — using CPU")

# ──────────────────────────────────────────────────────────────────────────────
# Configuration loader
# ──────────────────────────────────────────────────────────────────────────────

def _load_config() -> dict:
    cfg = {}
    for path in CONFIG_DIR.glob("*.yaml"):
        with open(path) as fh:
            cfg[path.stem] = yaml.safe_load(fh)
    return cfg


CONFIG = _load_config()

# ──────────────────────────────────────────────────────────────────────────────
# Data classes
# ──────────────────────────────────────────────────────────────────────────────

@dataclass
class RevisedSentence:
    sentence: str
    image_prompt: str
    highlighted: bool


@dataclass
class GeneratedIcon:
    sentence: str
    image_prompt: str
    highlighted: bool
    image_path: str

# ──────────────────────────────────────────────────────────────────────────────
# Services
# ──────────────────────────────────────────────────────────────────────────────

class GeminiDriver:
    """Thin wrapper around the Google Gemini API."""

    MODEL = "gemini-2.0-flash"  # stable public model; change to gemini-3-flash-preview if available

    def __init__(self):
        api_key = os.getenv("GOOGLE_API_KEY") or os.getenv("GEMINI_API_KEY")
        if not api_key:
            raise EnvironmentError(
                "GOOGLE_API_KEY (or GEMINI_API_KEY) is not set. "
                "Add it as a Space secret in the HuggingFace settings."
            )
        self.client = genai.Client(api_key=api_key)

    def generate_text(self, prompt: str) -> str:
        response = self.client.models.generate_content(
            model=self.MODEL, contents=prompt
        )
        return response.text


class GlobalSymbolsService:
    """Fetch AAC pictograms from the Global Symbols public API."""

    BASE_URL = "https://globalsymbols.com/api/v1/labels/search"

    def __init__(self, timeout: int = 10):
        self.timeout = timeout
        self.session = requests.Session()

    def search_and_download(
        self,
        query: str,
        output_path: Path,
        symbolset: str = "arasaac",
        language: str = "eng",
    ) -> Optional[Path]:
        try:
            resp = self.session.get(
                self.BASE_URL,
                params={
                    "query": query,
                    "symbolset": symbolset,
                    "language": language,
                    "limit": 1,
                },
                timeout=self.timeout,
            )
            resp.raise_for_status()
            results = resp.json()
        except requests.RequestException as exc:
            logger.warning(f"Global Symbols API error for '{query}': {exc}")
            return None

        if not results:
            logger.info(f"No symbol found for '{query}'")
            return None

        image_url = results[0].get("picto", {}).get("image_url")
        if not image_url:
            return None

        try:
            img_resp = self.session.get(image_url, timeout=self.timeout)
            img_resp.raise_for_status()
            output_path = Path(output_path)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            output_path.write_bytes(img_resp.content)
            return output_path
        except requests.RequestException as exc:
            logger.warning(f"Failed to download symbol image: {exc}")
            return None


class IconGenerator:
    """
    Generate EasyRead-style pictogram icons using Stable Diffusion v1.5 + LoRA.

    The generator is loaded lazily on first use so the Space starts up quickly
    even if GPU initialisation takes a few seconds.
    """

    BASE_MODEL = "runwayml/stable-diffusion-v1-5"
    NEGATIVE_PROMPT = (
        "blurry, photo, photograph, realistic, complex, detailed background"
    )
    INSTANCE_TOKEN = "sks"

    def __init__(self):
        from diffusers import StableDiffusionPipeline, DPMSolverMultistepScheduler
        from peft import PeftModel

        logger.info(f"Loading Stable Diffusion pipeline on {DEVICE} …")
        self.pipeline = StableDiffusionPipeline.from_pretrained(
            self.BASE_MODEL,
            torch_dtype=TORCH_DTYPE,
            safety_checker=None,
        )
        self.pipeline.scheduler = DPMSolverMultistepScheduler.from_config(
            self.pipeline.scheduler.config
        )

        weights_available = (
            WEIGHTS_DIR.exists() and any(WEIGHTS_DIR.iterdir())
        )
        if weights_available:
            logger.info(f"Loading LoRA weights from {WEIGHTS_DIR}")
            self.pipeline.unet = PeftModel.from_pretrained(
                self.pipeline.unet, str(WEIGHTS_DIR)
            )
        else:
            logger.info("No LoRA weights found — using base SD v1.5 model")

        self.pipeline = self.pipeline.to(DEVICE)
        self.pipeline.enable_attention_slicing()

        if DEVICE == "cuda":
            try:
                self.pipeline.enable_xformers_memory_efficient_attention()
                logger.info("xformers memory-efficient attention enabled")
            except Exception:
                logger.info("xformers not available; using default attention")

        logger.info("Stable Diffusion model ready")

    def generate(self, prompt: str, seed: Optional[int] = None) -> Image.Image:
        full_prompt = f"{self.INSTANCE_TOKEN} {prompt}"
        generator = None
        if seed is not None:
            generator = torch.Generator(device=DEVICE).manual_seed(seed)

        with torch.inference_mode():
            result = self.pipeline(
                prompt=full_prompt,
                negative_prompt=self.NEGATIVE_PROMPT,
                num_inference_steps=30,
                guidance_scale=7.5,
                height=256,
                width=256,
                generator=generator,
            )
        return result.images[0]


# ──────────────────────────────────────────────────────────────────────────────
# Lazy singletons
# ──────────────────────────────────────────────────────────────────────────────

_gemini: Optional[GeminiDriver] = None
_icon_generator: Optional[IconGenerator] = None
_global_symbols = GlobalSymbolsService()


def _get_gemini() -> GeminiDriver:
    global _gemini
    if _gemini is None:
        _gemini = GeminiDriver()
    return _gemini


def _get_icon_generator() -> IconGenerator:
    global _icon_generator
    if _icon_generator is None:
        _icon_generator = IconGenerator()
    return _icon_generator


# ──────────────────────────────────────────────────────────────────────────────
# Core AI pipeline
# ──────────────────────────────────────────────────────────────────────────────

def _parse_json_response(raw: str) -> dict:
    """Strip optional markdown fences and parse JSON."""
    cleaned = raw.strip()
    if cleaned.startswith("```"):
        # Remove opening fence (```json or ```)
        cleaned = cleaned.split("\n", 1)[1]
        # Remove closing fence
        cleaned = cleaned.rsplit("```", 1)[0]
    return json.loads(cleaned.strip())


def simplify_text(text: str, context: str = "", terms: str = "") -> dict:
    template = CONFIG["simplify_text"]["system_message"]
    extra = ""
    if context.strip():
        extra += f"\n\n# Additional Context\n{context.strip()}"
    if terms.strip():
        extra += f"\n\n# Terms to Preserve (do not simplify these)\n{terms.strip()}"
    prompt = f"{template}{extra}\n\nBellow is the Input Text to simplify:\n\n{text}\n\n"
    raw = _get_gemini().generate_text(prompt)
    try:
        return _parse_json_response(raw)
    except json.JSONDecodeError:
        logger.error(f"simplify_text: JSON parse failed. Raw:\n{raw}")
        return {"error": "JSON parse error", "raw_response": raw}


def validate_text(original: str, simplified: list) -> dict:
    template = CONFIG["validate_text"]["system_message"]
    user_msg = CONFIG["validate_text"]["user_message_template"].format(
        original_markdown=original,
        simplified_sentences=json.dumps(simplified),
    )
    raw = _get_gemini().generate_text(template + "\n" + user_msg)
    try:
        return _parse_json_response(raw)
    except json.JSONDecodeError:
        logger.error(f"validate_text: JSON parse failed. Raw:\n{raw}")
        return {"missing_info": "", "extra_info": "", "other_feedback": ""}


def revise_text(original: str, simplified: list, feedback: dict) -> dict:
    template = CONFIG["revise_text"]["system_message"]
    user_msg = CONFIG["revise_text"]["user_message_template"].format(
        original_markdown=original,
        simplified_sentences=json.dumps(simplified),
        validation_feedback=json.dumps(feedback),
    )
    raw = _get_gemini().generate_text(template + "\n" + user_msg)
    try:
        return _parse_json_response(raw)
    except json.JSONDecodeError:
        logger.error(f"revise_text: JSON parse failed. Raw:\n{raw}")
        return {"revised_sentences": simplified}


def generate_icons(sentences: list) -> dict:
    """
    For each sentence: try Global Symbols API first, fall back to local SD model.
    Images are saved to a temp directory keyed by a unique request ID.
    """
    request_id = str(uuid4())
    request_dir = TEMP_DIR / request_id
    request_dir.mkdir(parents=True, exist_ok=True)

    for sentence in sentences:
        prompt = sentence["image_prompt"]
        # Make a filesystem-safe filename (max 60 chars)
        safe_name = "_".join(prompt.split())[:60]
        image_path = request_dir / f"{safe_name}.png"

        # 1. Try Global Symbols
        downloaded = _global_symbols.search_and_download(
            query=prompt, output_path=image_path
        )

        # 2. Fall back to local Stable Diffusion
        if not downloaded:
            logger.info(f"Falling back to local SD for: '{prompt}'")
            try:
                img = _get_icon_generator().generate(prompt)
                img.save(image_path)
            except Exception as exc:
                logger.error(f"Local icon generation failed for '{prompt}': {exc}")
                continue  # skip this sentence's icon

        sentence["image_path"] = str(image_path)

    return {"request_id": request_id, "icons": sentences}


# ──────────────────────────────────────────────────────────────────────────────
# Export utilities
# ──────────────────────────────────────────────────────────────────────────────

def _export_docx_bytes(title: str, sentences: list, image_data: dict) -> bytes:
    doc = Document()
    title_p = doc.add_heading(title, level=0)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    for item in sentences:
        sentence = item.get("sentence", "")
        highlighted = item.get("highlighted", False)
        image_path = item.get("image_path", "")

        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        row = table.rows[0]
        icon_cell, text_cell = row.cells[0], row.cells[1]

        img_key = Path(image_path).name if image_path else None
        if img_key and img_key in image_data:
            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                tmp.write(image_data[img_key])
                tmp_path = tmp.name
            try:
                p = icon_cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(tmp_path, width=Inches(1.2))
            finally:
                if os.path.exists(tmp_path):
                    os.remove(tmp_path)
        else:
            r = icon_cell.paragraphs[0].add_run("[Icon]")
            r.font.color.rgb = RGBColor(128, 128, 128)

        p = text_cell.paragraphs[0]
        run = p.add_run(sentence)
        run.font.size = Pt(14)
        if highlighted:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 102, 204)

        doc.add_paragraph()

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        doc.save(tmp.name)
        tmp_path = tmp.name
    with open(tmp_path, "rb") as fh:
        data = fh.read()
    os.remove(tmp_path)
    return data


def _export_markdown_bytes(title: str, sentences: list) -> bytes:
    lines = [f"# {title}", "", "---", ""]
    for item in sentences:
        sentence = item.get("sentence", "")
        highlighted = item.get("highlighted", False)
        lines.append(f"**{sentence}**" if highlighted else sentence)
        lines.append("")
    lines += ["---", "", "*Generated by EasyRead*"]
    return "\n".join(lines).encode("utf-8")


# ──────────────────────────────────────────────────────────────────────────────
# App state  (single-user; for multi-user use gr.State instead)
# ──────────────────────────────────────────────────────────────────────────────

_state: dict = {
    "title": "",
    "revised_sentences": [],
    "request_id": "",
    "icons": [],
    "image_data": {},
}


# ──────────────────────────────────────────────────────────────────────────────
# Gradio event handlers
# ──────────────────────────────────────────────────────────────────────────────

def handle_simplify(text: str, context: str, terms: str):
    if not text.strip():
        return (
            gr.update(visible=True, value="Please enter some text to simplify."),
            gr.update(visible=False),
            gr.update(visible=False),
            "", "", None,
        )

    try:
        simple = simplify_text(text, context, terms)
        if "error" in simple:
            raise ValueError(simple.get("raw_response", "Unknown error"))

        validation = validate_text(text, simple["simplified_sentences"])
        revision = revise_text(text, simple["simplified_sentences"], validation)

        _state["title"] = simple.get("title", "Easy Read Document")
        _state["revised_sentences"] = [
            RevisedSentence(**s) for s in revision.get("revised_sentences", [])
        ]

        df_data = [
            [i + 1, s.sentence, s.image_prompt, s.highlighted]
            for i, s in enumerate(_state["revised_sentences"])
        ]
        val_text = (
            f"**Missing Information:** {validation.get('missing_info', 'None')}\n\n"
            f"**Extra Information:** {validation.get('extra_info', 'None')}\n\n"
            f"**Other Feedback:** {validation.get('other_feedback', 'None')}"
        )

        return (
            gr.update(visible=False),
            gr.update(visible=True),
            gr.update(visible=False),
            f"## {_state['title']}",
            val_text,
            df_data,
        )

    except Exception as exc:
        logger.error(f"handle_simplify error: {exc}", exc_info=True)
        msg = str(exc)
        if any(x in msg.lower() for x in ["api key", "environment"]):
            user_msg = f"Configuration error: {msg}"
        else:
            user_msg = f"Something went wrong: {msg[:300]}"
        return (
            gr.update(visible=True, value=user_msg),
            gr.update(visible=False),
            gr.update(visible=False),
            "", "", None,
        )


def handle_generate_icons(table_data):
    # Parse the editable dataframe
    rows = []
    if hasattr(table_data, "values"):
        rows = table_data.values.tolist()
    elif isinstance(table_data, dict) and "data" in table_data:
        rows = table_data["data"]
    elif isinstance(table_data, list):
        rows = table_data

    sentences = []
    for row in rows:
        if not isinstance(row, (list, tuple)) or len(row) < 4:
            continue
        sentences.append({
            "sentence": str(row[1]),
            "image_prompt": str(row[2]),
            "highlighted": bool(row[3]),
        })

    if not sentences:
        return (
            gr.update(visible=True, value="No sentences to process."),
            gr.update(visible=False),
            gr.update(visible=False),
            gr.update(visible=False),
            None, None,
        )

    try:
        result = generate_icons(sentences)
        _state["request_id"] = result["request_id"]
        _state["icons"] = result["icons"]
        _state["image_data"] = {}

        results_html = f'<div class="results-container"><h2>{_state["title"]}</h2>'

        for icon in result["icons"]:
            img_path = icon.get("image_path", "")
            img_src = ""
            if img_path and Path(img_path).exists():
                img_key = Path(img_path).name
                with open(img_path, "rb") as fh:
                    img_bytes = fh.read()
                _state["image_data"][img_key] = img_bytes
                img_b64 = base64.b64encode(img_bytes).decode("utf-8")
                img_src = f"data:image/png;base64,{img_b64}"

            hl_class = "highlighted" if icon.get("highlighted") else ""
            badge = (
                '<span class="badge">Key Point</span>'
                if icon.get("highlighted") else ""
            )
            img_tag = (
                f'<img src="{img_src}" alt="icon" />'
                if img_src
                else '<div class="placeholder">No icon</div>'
            )
            results_html += f"""
            <div class="sentence-row {hl_class}">
                <div class="icon-col">{img_tag}</div>
                <div class="text-col">{badge}<p>{icon["sentence"]}</p></div>
            </div>"""

        results_html += "</div>"

        return (
            gr.update(visible=False),
            gr.update(visible=True),
            gr.update(visible=True, value=results_html),
            gr.update(visible=True),
            None, None,
        )

    except Exception as exc:
        logger.error(f"handle_generate_icons error: {exc}", exc_info=True)
        return (
            gr.update(visible=True, value=f"Icon generation failed: {str(exc)[:300]}"),
            gr.update(visible=False),
            gr.update(visible=False),
            gr.update(visible=False),
            None, None,
        )


def handle_export_docx():
    if not _state["icons"]:
        return None
    data = _export_docx_bytes(_state["title"], _state["icons"], _state["image_data"])
    with tempfile.NamedTemporaryFile(suffix=".docx", prefix="easyread_", delete=False) as fh:
        fh.write(data)
        return fh.name


def handle_export_markdown():
    if not _state["icons"]:
        return None
    data = _export_markdown_bytes(_state["title"], _state["icons"])
    with tempfile.NamedTemporaryFile(suffix=".md", prefix="easyread_", delete=False) as fh:
        fh.write(data)
        return fh.name


def handle_reset():
    _state.update({
        "title": "", "revised_sentences": [],
        "request_id": "", "icons": [], "image_data": {},
    })
    return (
        "", "", "",
        gr.update(visible=False),
        gr.update(visible=False),
        gr.update(visible=False),
        gr.update(visible=False),
    )


# ──────────────────────────────────────────────────────────────────────────────
# Custom CSS  (identical to original frontend)
# ──────────────────────────────────────────────────────────────────────────────

CUSTOM_CSS = """
.section-box {
    border: 1px solid #e0e0e0;
    border-radius: 12px;
    padding: 20px;
    margin-top: 20px;
    background: #fafafa;
}
.results-container {
    max-width: 100%;
    font-family: Arial, sans-serif;
}
.results-container h2 {
    text-align: center;
    color: #333;
    margin-bottom: 25px;
    padding-bottom: 12px;
    border-bottom: 2px solid #0066cc;
}
.sentence-row {
    display: flex;
    align-items: center;
    padding: 15px;
    margin-bottom: 12px;
    background: white;
    border-radius: 10px;
    border-left: 4px solid #ddd;
    box-shadow: 0 1px 3px rgba(0,0,0,0.08);
}
.sentence-row.highlighted {
    background: #e6f3ff;
    border-left: 4px solid #0066cc;
}
.icon-col {
    flex: 0 0 100px;
    text-align: center;
    margin-right: 20px;
}
.icon-col img {
    max-width: 80px;
    max-height: 80px;
    border-radius: 8px;
    box-shadow: 0 2px 4px rgba(0,0,0,0.1);
}
.icon-col .placeholder {
    width: 80px;
    height: 80px;
    background: #eee;
    border-radius: 8px;
    display: flex;
    align-items: center;
    justify-content: center;
    color: #999;
    font-size: 12px;
}
.text-col { flex: 1; }
.text-col p {
    margin: 0;
    font-size: 16px;
    line-height: 1.5;
    color: #333;
}
.badge {
    display: inline-block;
    background: #0066cc;
    color: white;
    padding: 2px 8px;
    border-radius: 10px;
    font-size: 11px;
    margin-bottom: 6px;
}
.error-box {
    background: #fee;
    border: 1px solid #fcc;
    color: #c00;
    padding: 12px;
    border-radius: 8px;
    margin-bottom: 15px;
}
.loading-box {
    background: #e8f4fd;
    border: 1px solid #b3d9f7;
    color: #0066cc;
    padding: 16px;
    border-radius: 8px;
    margin: 15px 0;
    text-align: center;
    font-weight: 500;
}
"""

# ──────────────────────────────────────────────────────────────────────────────
# Gradio UI
# ──────────────────────────────────────────────────────────────────────────────

with gr.Blocks(
    title="EasyRead - Document Simplification",
    theme=gr.themes.Soft(),
    css=CUSTOM_CSS,
) as demo:

    gr.Markdown(
        """
        # EasyRead Document Simplification
        Convert complex documents into Easy Read format with pictogram icons.

        > Set **`GOOGLE_API_KEY`** as a Space secret to enable AI simplification.
        """
    )

    # ── Step 1: Input ──────────────────────────────────────────────────────────
    with gr.Group():
        gr.Markdown("### Step 1: Enter Your Text")
        input_text = gr.Textbox(
            label="Text to Simplify",
            placeholder="Paste or type the text you want to convert to Easy Read format…",
            lines=8,
        )
        with gr.Row():
            context_input = gr.Textbox(
                label="Custom Context (Optional)",
                placeholder="Add context to help with simplification…",
                lines=2,
                scale=1,
            )
            terms_input = gr.Textbox(
                label="Terms to Preserve (Optional)",
                placeholder="Comma-separated terms…",
                lines=2,
                scale=1,
            )
        with gr.Row():
            simplify_btn = gr.Button("Simplify Text", variant="primary", size="lg")
            reset_btn = gr.Button("Reset", variant="secondary")

    error_box = gr.Markdown(visible=False, elem_classes=["error-box"])
    loading_box = gr.Markdown(visible=False, elem_classes=["loading-box"])

    # ── Step 2: Review ─────────────────────────────────────────────────────────
    with gr.Group(visible=False, elem_classes=["section-box"]) as review_section:
        gr.Markdown("### Step 2: Review & Approve Sentences")
        title_display = gr.Markdown()
        with gr.Accordion("Validation Feedback", open=False):
            validation_display = gr.Markdown()
        gr.Markdown(
            "Edit sentences or image prompts below if needed, "
            "then click **Generate Icons**."
        )
        sentences_table = gr.Dataframe(
            headers=["#", "Sentence", "Image Prompt", "Highlighted"],
            datatype=["number", "str", "str", "bool"],
            col_count=(4, "fixed"),
            interactive=True,
            wrap=True,
        )
        generate_btn = gr.Button("Generate Icons", variant="primary", size="lg")

    # ── Step 3: Results ────────────────────────────────────────────────────────
    with gr.Group(visible=False, elem_classes=["section-box"]) as results_section:
        gr.Markdown("### Step 3: Your Easy Read Document")
        results_display = gr.HTML()
        with gr.Row(visible=False) as export_row:
            docx_btn = gr.Button("Export to Word (.docx)", variant="secondary")
            md_btn = gr.Button("Export to Markdown (.md)", variant="secondary")
        with gr.Row():
            docx_download = gr.File(label="Word Document", visible=False)
            md_download = gr.File(label="Markdown File", visible=False)

    # ── Event wiring ───────────────────────────────────────────────────────────
    simplify_btn.click(
        fn=lambda: gr.update(
            visible=True, value="Simplifying your text… This may take a moment."
        ),
        outputs=[loading_box],
    ).then(
        fn=handle_simplify,
        inputs=[input_text, context_input, terms_input],
        outputs=[
            error_box, review_section, results_section,
            title_display, validation_display, sentences_table,
        ],
    ).then(
        fn=lambda: gr.update(visible=False),
        outputs=[loading_box],
    )

    generate_btn.click(
        fn=lambda: gr.update(
            visible=True, value="Generating icons… This may take a moment."
        ),
        outputs=[loading_box],
    ).then(
        fn=handle_generate_icons,
        inputs=[sentences_table],
        outputs=[
            error_box, results_section, results_display,
            export_row, docx_download, md_download,
        ],
    ).then(
        fn=lambda: gr.update(visible=False),
        outputs=[loading_box],
    )

    docx_btn.click(fn=handle_export_docx, outputs=[docx_download]).then(
        fn=lambda: gr.update(visible=True), outputs=[docx_download]
    )
    md_btn.click(fn=handle_export_markdown, outputs=[md_download]).then(
        fn=lambda: gr.update(visible=True), outputs=[md_download]
    )
    reset_btn.click(
        fn=handle_reset,
        outputs=[
            input_text, context_input, terms_input,
            error_box, loading_box, review_section, results_section,
        ],
    )


# ──────────────────────────────────────────────────────────────────────────────
# Entry point
# ──────────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    demo.launch()
